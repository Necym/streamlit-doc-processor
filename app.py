import streamlit as st
import pandas as pd
from docx import Document
import re
from io import BytesIO

# Function to extract the question prompt, answers, and explanation from the given text
def extract_prompt_answers_and_explanation(row):
    question_text = row['Question']
    explanation = row['Explanation']
    
    # Split the text into prompt and answers
    parts = re.split(r'(?=\bA\.)', question_text, maxsplit=1)
    prompt = parts[0].strip()
    
    # Remove the labels (A., B., C., D.) from the answers
    answers = []
    if len(parts) > 1:
        raw_answers = re.split(r'\n(?=[A-Z]\.)', parts[1].strip())
        for answer in raw_answers:
            # Remove the answer label (e.g., A., B., C., D.)
            clean_answer = re.sub(r'^[A-Z]\.\s*', '', answer).strip()
            answers.append(clean_answer)
    
    return prompt, answers, explanation


# Function for Version A (original process)
def scan_word_document_version_a(word_file, excel_file, question_limit):
    df = pd.read_excel(excel_file, sheet_name='Simulated')
    doc = Document(word_file)
    required_columns = ["ID", "Type", "Source Text", "Translation"]

    def column_match(headers, required_columns):
        return all(any(required.lower() in header.lower() for required in required_columns) for header in headers)

    questions_processed = 0
    for table in doc.tables:
        first_row = table.rows[0]
        headers = [cell.text.strip() for cell in first_row.cells]

        if column_match(headers, required_columns):
            for i in range(1, len(table.rows)):  # Skip the header row
                row = table.rows[i]
                if len(row.cells) < 4:
                    continue

                id_value = row.cells[0].text.strip()
                type_value = row.cells[1].text.strip()
                source_text_value = row.cells[2].text.strip()

                if type_value.lower() == "question number" and re.search(r'Question\s*(\d+)', source_text_value, re.IGNORECASE):
                    match = re.search(r'Question\s*(\d+)', source_text_value, re.IGNORECASE)
                    question_number = int(match.group(1))

                    excel_row_index = question_number

                    if excel_row_index <= len(df):
                        excel_prompt, excel_answers, excel_explanation = extract_prompt_answers_and_explanation(df.iloc[excel_row_index - 1])

                        if i + 2 < len(table.rows):
                            prompt_row = table.rows[i + 2]
                            prompt_cell = prompt_row.cells[3]
                            prompt_cell.text = excel_prompt

                            for j, answer in enumerate(excel_answers):
                                if i + 3 + j < len(table.rows):
                                    answer_row = table.rows[i + 3 + j]
                                    answer_cell = answer_row.cells[3]
                                    answer_cell.text = answer

                            explanation_row_index = i + 3 + len(excel_answers) + 1
                            if explanation_row_index < len(table.rows):
                                explanation_row = table.rows[explanation_row_index]
                                explanation_cell = explanation_row.cells[3]
                                explanation_cell.text = excel_explanation
                    else:
                        st.write(f"No matching row found in Excel for Question {question_number}.")

                    questions_processed += 1
                    if questions_processed >= question_limit:
                        output_buffer = BytesIO()
                        doc.save(output_buffer)
                        output_buffer.seek(0)
                        return output_buffer, "Processing complete."
        else:
            st.write(f"Table skipped due to missing required columns: {headers}")

    output_buffer = BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    return output_buffer, "Processing complete."


# Function for Version B (new process)
def scan_word_document_version_b(word_file, excel_file, question_limit):
    df = pd.read_excel(excel_file, sheet_name='Simulated')
    doc = Document(word_file)
    required_columns = ["ID", "Type", "Source Text", "Translation"]

    def column_match(headers, required_columns):
        return all(any(required.lower() in header.lower() for required in required_columns) for header in headers)

    questions_processed = 0
    for table in doc.tables:
        first_row = table.rows[0]
        headers = [cell.text.strip() for cell in first_row.cells]

        if column_match(headers, required_columns):
            for i in range(1, len(table.rows)):  # Skip the header row
                row = table.rows[i]
                if len(row.cells) < 4:
                    continue

                id_value = row.cells[0].text.strip()
                type_value = row.cells[1].text.strip()
                source_text_value = row.cells[2].text.strip()

                if type_value.lower() == "slide name" and re.search(r'Question\s*(\d+)', source_text_value, re.IGNORECASE):
                    match = re.search(r'Question\s*(\d+)', source_text_value, re.IGNORECASE)
                    question_number = int(match.group(1))

                    excel_row_index = question_number

                    if excel_row_index <= len(df):
                        excel_prompt, excel_answers, excel_explanation = extract_prompt_answers_and_explanation(df.iloc[excel_row_index - 1])

                        if i + 2 < len(table.rows):
                            prompt_row = table.rows[i + 2]
                            prompt_cell = prompt_row.cells[3]
                            prompt_cell.text = excel_prompt

                            for j, answer in enumerate(excel_answers):
                                if i + 3 + j < len(table.rows):
                                    answer_row = table.rows[i + 3 + j]
                                    answer_cell = answer_row.cells[3]
                                    answer_cell.text = answer

                            explanation_row_index = i + 3 + len(excel_answers)
                            if explanation_row_index < len(table.rows):
                                explanation_row = table.rows[explanation_row_index]
                                explanation_cell = explanation_row.cells[3]
                                explanation_cell.text = excel_explanation
                    else:
                        st.write(f"No matching row found in Excel for Question {question_number}.")

                    questions_processed += 1
                    if questions_processed >= question_limit:
                        output_buffer = BytesIO()
                        doc.save(output_buffer)
                        output_buffer.seek(0)
                        return output_buffer, "Processing complete."
        else:
            st.write(f"Table skipped due to missing required columns: {headers}")

    output_buffer = BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    return output_buffer, "Processing complete."


# Streamlit app logic
st.title("Document Processor")

# Add a selectbox to choose between Version A and Version B
version_choice = st.selectbox("Select Version", ["Version A", "Version B"])

word_file = st.file_uploader("Upload Word Document", type=["docx"])
excel_file = st.file_uploader("Upload Excel Document", type=["xlsx"])

if word_file and excel_file:
    question_limit = st.number_input("How many questions would you like to process?", min_value=1, value=5)
    if st.button("Process"):
        if version_choice == "Version A":
            output_buffer, output_message = scan_word_document_version_a(BytesIO(word_file.read()), BytesIO(excel_file.read()), question_limit)
        else:
            output_buffer, output_message = scan_word_document_version_b(BytesIO(word_file.read()), BytesIO(excel_file.read()), question_limit)
        
        st.write(output_message)
        st.success("Processing complete. Download the updated Word document below.")
        st.download_button(
            label="Download updated document",
            data=output_buffer,
            file_name="updated_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
